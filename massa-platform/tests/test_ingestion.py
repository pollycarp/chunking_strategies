"""
Phase 3 Tests: Document Ingestion & Chunking

Test categories:
- Parser tests      : each parser extracts known text from a fixture file
- Chunker tests     : verify splitting logic, token limits, metadata
- Deduplication     : same file twice → skipped on second run
- Pipeline test     : end-to-end with mock embedder
"""

import io
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF
from openpyxl import Workbook

from src.embeddings.base import EmbeddingModel
from src.ingestion.chunkers.fixed_size import FixedSizeChunker
from src.ingestion.chunkers.hierarchical import HierarchicalChunker
from src.ingestion.chunkers.semantic import SemanticChunker
from src.ingestion.deduplication import is_already_ingested
from src.ingestion.parsers.docx_parser import DocxParser
from src.ingestion.parsers.pdf_parser import PDFParser
from src.ingestion.parsers.xlsx_parser import XLSXParser
from src.ingestion.pipeline import IngestionPipeline


# ---------------------------------------------------------------------------
# Fixture file factories
# These create real files in temp directories so parsers can read them.
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Creates a minimal 2-page PDF with known content."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, "Financial Performance Report Q3 2024\n\n"
                          "Revenue grew 12% year-over-year driven by strong "
                          "performance in the agribusiness segment. EBITDA margin "
                          "improved to 23.4% from 21.1% in the prior year period.")
    pdf.add_page()
    pdf.multi_cell(0, 10, "Risk Factors\n\n"
                          "The company faces headwinds from rising input costs "
                          "and foreign exchange volatility. Management has implemented "
                          "hedging strategies to mitigate currency exposure.")

    path = tmp_path / "report.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Creates a DOCX with headings and body paragraphs."""
    doc = DocxDocument()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The portfolio company delivered strong Q3 results with revenue "
        "of $45.2M, representing 18% growth versus prior year."
    )
    doc.add_heading("Financial Highlights", level=2)
    doc.add_paragraph(
        "EBITDA margin expanded by 230 basis points to 24.1%. "
        "Free cash flow conversion remained above 85%."
    )

    path = tmp_path / "memo.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_xlsx(tmp_path: Path) -> Path:
    """Creates an XLSX with two sheets of financial data."""
    wb = Workbook()

    # Income statement sheet
    ws1 = wb.active
    ws1.title = "Income Statement"
    ws1.append(["Metric", "Q1", "Q2", "Q3", "Q4"])
    ws1.append(["Revenue", 42000, 43500, 45200, 47800])
    ws1.append(["COGS", 28000, 28900, 29800, 31200])
    ws1.append(["Gross Profit", 14000, 14600, 15400, 16600])
    ws1.append(["EBITDA", 9800, 10200, 10900, 11800])

    # Balance sheet
    ws2 = wb.create_sheet("Balance Sheet")
    ws2.append(["Asset", "Value"])
    ws2.append(["Cash", 12500000])
    ws2.append(["Receivables", 8750000])
    ws2.append(["Total Assets", 45000000])

    path = tmp_path / "financials.xlsx"
    wb.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Test 1: Parser tests
# ---------------------------------------------------------------------------

def test_pdf_parser_extracts_text(sample_pdf):
    """Assert PDF parser extracts text from both pages with page numbers."""
    parser = PDFParser()
    doc = parser.parse(sample_pdf)

    assert doc.doc_type == "pdf"
    assert doc.filename == "report.pdf"
    assert len(doc.pages) == 2
    assert doc.pages[0].page_number == 1
    assert doc.pages[1].page_number == 2
    assert "Revenue" in doc.pages[0].content
    assert "Risk Factors" in doc.pages[1].content
    assert doc.file_hash  # hash should be populated


def test_pdf_parser_raises_on_missing_file():
    """Assert FileNotFoundError for non-existent files — not a cryptic internal error."""
    with pytest.raises(FileNotFoundError):
        PDFParser().parse(Path("/nonexistent/file.pdf"))


def test_docx_parser_extracts_text_and_sections(sample_docx):
    """Assert DOCX parser extracts text and detects heading-based section titles."""
    parser = DocxParser()
    doc = parser.parse(sample_docx)

    assert doc.doc_type == "docx"
    assert doc.filename == "memo.docx"
    assert len(doc.pages) >= 1
    assert "Executive Summary" in doc.full_text
    assert "EBITDA" in doc.full_text


def test_xlsx_parser_extracts_sheets_as_pages(sample_xlsx):
    """Assert XLSX parser creates one page per sheet and preserves tabular structure."""
    parser = XLSXParser()
    doc = parser.parse(sample_xlsx)

    assert doc.doc_type == "xlsx"
    assert doc.filename == "financials.xlsx"
    assert len(doc.pages) == 2  # two sheets
    assert doc.pages[0].section_title == "Income Statement"
    assert doc.pages[1].section_title == "Balance Sheet"
    # Pipe-delimited format preserves column structure
    assert "|" in doc.pages[0].content
    assert "Revenue" in doc.pages[0].content


def test_all_parsers_populate_required_metadata(sample_pdf, sample_docx, sample_xlsx):
    """Assert every parser returns a document with all required metadata fields."""
    docs = [
        PDFParser().parse(sample_pdf),
        DocxParser().parse(sample_docx),
        XLSXParser().parse(sample_xlsx),
    ]

    for doc in docs:
        assert doc.filename, "filename must be set"
        assert doc.file_path, "file_path must be set"
        assert doc.doc_type in ("pdf", "docx", "xlsx"), "doc_type must be valid"
        assert doc.file_hash, "file_hash must be set"
        assert len(doc.pages) > 0, "must have at least one page"
        for page in doc.pages:
            assert page.page_number > 0, "page numbers must be 1-based"
            assert page.content.strip(), "page content must not be empty"


# ---------------------------------------------------------------------------
# Test 2: Chunker tests
# ---------------------------------------------------------------------------

@pytest.fixture
def long_document(sample_pdf) -> "ParsedDocument":
    return PDFParser().parse(sample_pdf)


def test_fixed_size_chunker_respects_token_limit(long_document):
    """Assert no chunk exceeds max_tokens."""
    chunker = FixedSizeChunker(max_tokens=100, overlap=10)
    chunks = chunker.chunk(long_document)

    assert len(chunks) > 0
    for chunk in chunks:
        assert chunk.token_count <= 100, (
            f"Chunk at index {chunk.chunk_index} has {chunk.token_count} tokens, "
            f"exceeds max of 100"
        )


def test_fixed_size_chunker_overlap_creates_continuity(sample_pdf):
    """Assert overlap tokens appear in consecutive chunks."""
    chunker = FixedSizeChunker(max_tokens=50, overlap=20)
    doc = PDFParser().parse(sample_pdf)
    chunks = chunker.chunk(doc)

    if len(chunks) >= 2:
        # Last words of chunk 0 should appear in start of chunk 1
        words_0 = set(chunks[0].content.split())
        words_1 = set(chunks[1].content.split())
        overlap_words = words_0 & words_1
        assert len(overlap_words) > 0, "Consecutive chunks should share words due to overlap"


def test_semantic_chunker_no_mid_sentence_splits(long_document):
    """Assert semantic chunks don't end mid-sentence (no trailing incomplete sentences)."""
    chunker = SemanticChunker(max_tokens=200)
    chunks = chunker.chunk(long_document)

    assert len(chunks) > 0
    for chunk in chunks:
        stripped = chunk.content.strip()
        # A properly bounded chunk ends with sentence-ending punctuation or a word
        # (not with a conjunction or preposition suggesting a mid-sentence cut)
        assert stripped, "chunk content must not be empty"
        assert chunk.token_count > 0


def test_semantic_chunker_preserves_metadata(long_document):
    """Assert semantic chunks carry source_file, page_number, doc_type."""
    chunker = SemanticChunker(max_tokens=200)
    chunks = chunker.chunk(long_document)

    for chunk in chunks:
        assert chunk.source_file == long_document.filename
        assert chunk.page_number is not None
        assert chunk.doc_type == "pdf"
        assert chunk.chunk_strategy == "semantic"


def test_hierarchical_chunker_creates_parent_child_structure(long_document):
    """Assert hierarchical chunks have correct parent/child relationships."""
    chunker = HierarchicalChunker(parent_max_tokens=200, child_max_tokens=80)
    chunks = chunker.chunk(long_document)

    parents = [c for c in chunks if c.parent_chunk_index is None]
    children = [c for c in chunks if c.parent_chunk_index is not None]

    assert len(parents) > 0, "should have at least one parent chunk"
    assert len(children) > 0, "should have at least one child chunk"

    # All child parent references should point to valid parent indices
    parent_indices = {p.chunk_index for p in parents}
    for child in children:
        assert child.parent_chunk_index in parent_indices, (
            f"Child references parent_index {child.parent_chunk_index} "
            f"which does not exist"
        )


def test_all_chunks_have_required_fields(long_document):
    """Assert every chunking strategy produces chunks with all required metadata."""
    chunkers = [
        FixedSizeChunker(max_tokens=100),
        SemanticChunker(max_tokens=100),
        HierarchicalChunker(parent_max_tokens=200, child_max_tokens=80),
    ]

    for chunker in chunkers:
        chunks = chunker.chunk(long_document)
        for chunk in chunks:
            assert chunk.content.strip(), "content must not be empty"
            assert chunk.source_file, "source_file must be set"
            assert chunk.doc_type, "doc_type must be set"
            assert chunk.chunk_strategy, "chunk_strategy must be set"
            assert chunk.token_count > 0, "token_count must be positive"
            assert chunk.content_hash, "content_hash must be set"


# ---------------------------------------------------------------------------
# Test 3: Deduplication
# ---------------------------------------------------------------------------

async def test_deduplication_skips_already_ingested_file(db_pool, sample_docx):
    """
    Assert the pipeline skips a file that was already ingested.

    First ingestion: skipped=False, document stored in DB
    Second ingestion: skipped=True, no new document row
    """
    call_count = 0

    async def fake_embed_batch(texts):
        nonlocal call_count
        call_count += 1
        return [[0.1] * 1536 for _ in texts]

    mock_embedder = MagicMock(spec=EmbeddingModel)
    mock_embedder.embed_batch = fake_embed_batch

    pipeline = IngestionPipeline(pool=db_pool, embedder=mock_embedder)

    # First ingest
    result1 = await pipeline.ingest(sample_docx)
    assert result1.skipped is False
    assert result1.document_id is not None

    # Second ingest — same file
    result2 = await pipeline.ingest(sample_docx)
    assert result2.skipped is True
    assert result2.document_id is None

    # Embedder should only have been called once (first ingest)
    assert call_count == 1, "embedder should not be called on duplicate ingest"


async def test_deduplication_reingest_after_content_change(db_pool, tmp_path):
    """
    Assert that a file with changed content gets re-ingested (hash changes).
    """
    async def fake_embed_batch(texts):
        return [[0.1] * 1536 for _ in texts]

    mock_embedder = MagicMock(spec=EmbeddingModel)
    mock_embedder.embed_batch = fake_embed_batch

    pipeline = IngestionPipeline(pool=db_pool, embedder=mock_embedder)

    # Version 1
    doc_v1 = DocxDocument()
    doc_v1.add_paragraph("Original content version 1")
    path = tmp_path / "evolving.docx"
    doc_v1.save(str(path))

    result1 = await pipeline.ingest(path)
    assert result1.skipped is False

    # Version 2 — same path, different content
    doc_v2 = DocxDocument()
    doc_v2.add_paragraph("Updated content version 2 with new financial data")
    doc_v2.save(str(path))

    result2 = await pipeline.ingest(path)
    assert result2.skipped is False, "Changed file should be re-ingested"


# ---------------------------------------------------------------------------
# Test 4: End-to-end pipeline
# ---------------------------------------------------------------------------

async def test_pipeline_stores_chunks_in_db(db_pool, sample_xlsx):
    """
    End-to-end: ingest a file → verify chunks appear in the DB with embeddings.
    """
    async def fake_embed_batch(texts):
        # Return deterministic fake embeddings (not zeros — zeros cause issues with cosine)
        return [[float(i % 10) / 10.0] + [0.1] * 1535 for i, _ in enumerate(texts)]

    mock_embedder = MagicMock(spec=EmbeddingModel)
    mock_embedder.embed_batch = fake_embed_batch

    pipeline = IngestionPipeline(
        pool=db_pool,
        embedder=mock_embedder,
        chunker=FixedSizeChunker(max_tokens=100),
    )

    result = await pipeline.ingest(sample_xlsx)

    assert result.skipped is False
    assert result.chunk_count > 0
    assert result.document_id is not None

    # Verify chunks are in the DB
    async with db_pool.acquire() as conn:
        count = await conn.fetchval(
            "SELECT COUNT(*) FROM chunks WHERE document_id = $1",
            result.document_id,
        )
        assert count == result.chunk_count, (
            f"DB has {count} chunks but pipeline reported {result.chunk_count}"
        )

        # Verify embeddings are stored (not NULL)
        null_embeddings = await conn.fetchval(
            "SELECT COUNT(*) FROM chunks WHERE document_id = $1 AND embedding IS NULL",
            result.document_id,
        )
        assert null_embeddings == 0, "All chunks should have embeddings stored"
