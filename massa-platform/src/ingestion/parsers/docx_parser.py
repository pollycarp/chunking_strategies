from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.ingestion.models import ParsedDocument, ParsedPage
from src.ingestion.parsers.base import DocumentParser

# Word heading styles that indicate section boundaries
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}


class DocxParser(DocumentParser):
    """
    Extracts text from Word (.docx) files using python-docx.

    Word documents don't have pages — they have paragraphs.
    We simulate "pages" by grouping paragraphs into chunks of
    ~50 paragraphs, which approximates a page of content.

    We also extract section headings from paragraph styles
    (Heading 1, Heading 2, etc.) to populate section_title metadata.

    Why does section_title matter?
    In a 50-page investment memo, knowing a chunk came from the
    "Risk Factors" section vs "Financial Projections" section is
    critical context for the LLM — and for metadata filtering.
    """

    PARAGRAPHS_PER_PAGE = 50   # approximate grouping for page simulation

    def parse(self, file_path: Path) -> ParsedDocument:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")

        file_hash = self.compute_file_hash(file_path)
        doc = Document(str(file_path))

        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError(f"No text content found in {file_path.name}")

        pages: list[ParsedPage] = []
        current_section: str | None = None

        # Group paragraphs into simulated pages
        for page_num, start in enumerate(
            range(0, len(paragraphs), self.PARAGRAPHS_PER_PAGE), start=1
        ):
            page_paragraphs = paragraphs[start: start + self.PARAGRAPHS_PER_PAGE]

            # Track the last heading seen in this page for section_title
            section_in_page = current_section
            lines = []
            for para in page_paragraphs:
                if para.style.name in _HEADING_STYLES and para.text.strip():
                    current_section = para.text.strip()
                    section_in_page = current_section
                lines.append(para.text)

            content = "\n".join(lines).strip()
            if content:
                pages.append(ParsedPage(
                    content=content,
                    page_number=page_num,
                    section_title=section_in_page,
                ))

        return ParsedDocument(
            filename=file_path.name,
            file_path=str(file_path),
            doc_type="docx",
            pages=pages,
            file_hash=file_hash,
        )
